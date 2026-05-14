"""業務委託契約書 .md → .docx 変換スクリプト"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import sys
import re

HERE = os.path.dirname(os.path.abspath(__file__))


def set_cell_bg(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_default_font(doc: Document, font_name: str = "游明朝") -> None:
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0)


def add_list_item(doc: Document, text: str, indent_level: int = 0) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(0.6 * (indent_level + 1))


def build_contract_doc() -> None:
    doc = Document()
    set_default_font(doc)

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    add_title(doc, "業務委託契約書")

    intro = (
        "株式会社91&Co.（以下「甲」という）と株式会社THINGx（以下「乙」という）は、"
        "甲が乙に委託する業務に関して、以下のとおり業務委託契約（以下「本契約」という）を締結する。"
    )
    add_body(doc, intro)

    clauses = [
        {
            "heading": "第1条（目的）",
            "items": [
                "本契約は、甲が運営するインターネットプラットフォーム「MyFocus」（以下「本サービス」という）に関する開発・運用等の業務を、乙に委託するにあたっての、甲乙間の権利義務関係を定めることを目的とする。",
            ],
        },
        {
            "heading": "第2条（委託業務の内容）",
            "items": [
                "1. 甲は乙に対し、以下の業務（以下「本件業務」という）を委託し、乙はこれを受託する。",
                "　(1) 本サービスに関するソフトウェア開発業務（新機能追加・改修を含む）",
                "　(2) 本サービスの保守・運用業務（障害対応・軽微な不具合修正を含む）",
                "　(3) 本サービスに利用するインフラ（クラウドサービス等）の選定・設定・管理業務",
                "　(4) 前各号に付随する業務",
                "2. 本件業務の詳細、納期、成果物の仕様等については、甲乙協議のうえ別途個別の発注書または仕様書（以下「個別発注書」という）を作成し、これを本契約の一部とする。",
            ],
        },
        {
            "heading": "第3条（契約形態）",
            "items": [
                "1. 本契約は、民法上の準委任契約（民法656条）とし、乙は委任の本旨に従い、善良な管理者の注意をもって本件業務を遂行する。",
                "2. 本サービスの運営主体は甲であり、乙は本サービスの運営者ではない。本サービスに関する利用規約の当事者、特定商取引法に基づく表記上の販売事業者、決済事業者との契約主体、個人情報保護法上の個人情報取扱事業者は、いずれも甲とする。",
                "3. 本契約は、甲乙間に雇用関係・代理関係・共同事業関係を生じさせるものではない。",
            ],
        },
        {
            "heading": "第4条（委託料）",
            "items": [
                "1. 本件業務のうち、以下に掲げる継続的な保守・運用業務（以下「本件保守業務」という）の委託料は、月額0円とする。",
                "　(1) 本サービスの稼働監視",
                "　(2) 障害発生時の初動対応",
                "　(3) 軽微な不具合修正および軽微な設定変更",
                "　(4) 次条に定めるホスティング環境の提供",
                "2. 前項にいう「軽微な」作業とは、1件あたりの作業時間が概ね2時間以内で完了する範囲のものを指し、これに該当するか否かは甲乙協議のうえ決定する。",
                "3. 第1項に定める本件保守業務の範囲を超える開発業務（新機能追加・大規模改修・UIデザインの大幅変更・データ移行等、以下「追加開発業務」という）については、乙は甲に対して事前に見積書を提示し、甲乙協議のうえ個別発注書を作成して委託料を合意する。",
                "4. 追加開発業務の委託料は、個別発注書に定める支払条件に従い、甲が乙指定の銀行口座に振込む。振込手数料は甲の負担とする。",
            ],
        },
        {
            "heading": "第5条（ホスティング環境の提供）",
            "items": [
                "1. 乙は、本サービスの技術的実行環境として、自己の名義および費用において以下のクラウドサービス等（以下「本件ホスティング環境」という）を契約・維持し、甲に対して提供する。",
                "　(1) アプリケーションホスティング（Vercel 等）",
                "　(2) データベース・ストレージ・認証基盤（Supabase 等）",
                "　(3) メール配信基盤（Resend 等）",
                "　(4) その他、本サービスの稼働に必要と甲乙が合意した第三者サービス",
                "2. 前項の契約関係は、乙が本件業務を甲に提供するための技術的インフラとしての位置づけであり、乙は本件ホスティング環境の契約当事者となるが、本サービスそのものの運営者ではない。本サービス上で取り扱われるデータ（ユーザー・クリエイターの個人情報、コンテンツ、取引記録等、以下「本サービスデータ」という）の管理責任は、次条および第11条の定めに従い、すべて甲が負う。",
                "3. 本件ホスティング環境の提供は、第4条第1項第4号により月額0円の保守料に含まれるものとし、乙は甲に対して別途の対価を請求しない。ただし、以下のいずれかに該当する場合は、乙は甲に対して追加費用を請求できる。",
                "　(1) 本サービスの利用量・データ量の急増等により、クラウドサービス利用料が乙の負担可能な範囲（以下「負担上限」という）を超える場合",
                "　(2) AI コンテンツモデレーション等、本件ホスティング環境と性質を異にする有償サービスの導入が必要となる場合",
                "　(3) 甲の指示により、通常より高いプラン・リージョン・機能を利用する必要が生じた場合",
                "4. 前項の負担上限は、本契約締結時点において月額7,000円（消費税別）を目安とし、これを継続的に超過する見込みが生じた場合、乙は甲に対して速やかに通知し、甲乙協議のうえ費用分担または料金プラン変更等の対応を決定する。",
                "5. 決済事業者（Stripe 等）との契約関係については、本条の定めにかかわらず、甲が自ら契約当事者となり、甲名義の銀行口座を振込先として設定する。乙は、甲による当該契約の技術実装・設定作業を支援する。",
                "6. 乙は、自己の事業環境の変化、甲乙間もしくは乙と甲の関係会社との他の取引関係の変更・終了、その他相当の事由により、本件ホスティング環境を無償で提供することが困難となった場合、3ヶ月前の書面通知をもって、以下のいずれかの措置を甲に求めることができる。",
                "　(1) 本件ホスティング環境の有償化（クラウドサービス利用料実費＋消費税を甲が負担する運用への切替）",
                "　(2) 本件ホスティング環境の甲名義または甲が指定する第三者名義への移管（第16条に準じた移管措置）",
                "　(3) 本件ホスティング環境の提供の停止",
                "7. 前項の通知を受けた甲は、通知受領後30日以内に上記(1)ないし(3)のいずれを選択するかを乙に通知する。甲が期間内に選択を通知しない場合は、(2)を選択したものとみなす。",
            ],
        },
        {
            "heading": "第6条（成果物の納品・検収）",
            "items": [
                "1. 個別発注書に基づく成果物は、乙が個別発注書所定の方法（Git リポジトリへのコミット、本番環境へのデプロイ、その他甲乙が合意する方法）により甲に納品する。",
                "2. 甲は、成果物の納品後10営業日以内に検収を行い、合格した場合は乙に検収合格の通知を行う。",
                "3. 前項の期間内に甲から不合格の通知がない場合、成果物は検収合格したものとみなす。",
                "4. 乙は、検収合格後1ヶ月間に限り、個別発注書の仕様書と明らかに異なる不具合（以下「契約不適合」という）について、無償で修補する義務を負う。期間経過後の修補は、第4条または別途合意する委託料の対象とする。",
            ],
        },
        {
            "heading": "第7条（知的財産権の帰属）",
            "items": [
                "1. 乙が本件業務の遂行過程で作成した成果物（プログラム、ソースコード、設計書、デザイン等）の著作権（著作権法第27条および第28条の権利を含む）は、検収合格および当該成果物に対応する委託料の全額支払い完了の時点で、乙から甲へ移転する。",
                "2. 前項の規定にかかわらず、以下は乙の著作権または権利に帰属する。",
                "　(1) 本契約締結前から乙が有していた著作物・ノウハウ",
                "　(2) 汎用的に利用可能なモジュール・ライブラリ・テンプレート類",
                "3. 乙は、甲が成果物を本サービスの運営目的で無償かつ期間無制限で利用することを許諾する。",
                "4. 乙は、甲および本サービスのユーザーに対して、成果物について著作者人格権を行使しない。",
            ],
        },
        {
            "heading": "第8条（再委託）",
            "items": [
                "1. 乙は、本件業務の全部または一部を、甲の事前の書面による承諾を得たうえで、第三者に再委託することができる。",
                "2. 乙は、再委託先の行為について、自ら行ったものと同様の責任を負う。",
            ],
        },
        {
            "heading": "第9条（秘密保持）",
            "items": [
                "1. 甲および乙は、本契約の履行過程で知り得た相手方の技術上・営業上・運営上の情報（以下「秘密情報」という）を、相手方の事前の書面による承諾なく第三者に開示・漏洩してはならない。",
                "2. 前項の規定は、以下の情報には適用しない。",
                "　(1) 開示の時点で既に公知となっている情報",
                "　(2) 開示後、自己の責めに帰さない事由により公知となった情報",
                "　(3) 開示の時点で既に自己が保有していた情報",
                "　(4) 正当な権限を有する第三者から秘密保持義務を負うことなく入手した情報",
                "　(5) 裁判所・監督官庁等の法令に基づく開示請求に応じて開示する情報",
                "3. 本条の義務は、本契約終了後も3年間存続する。",
            ],
        },
        {
            "heading": "第10条（個人情報）",
            "items": [
                "1. 乙は、本件業務の遂行にあたり、甲から提供され、または本サービスのシステム上で取扱う個人情報について、個人情報の保護に関する法律その他関連法令を遵守する。",
                "2. 乙は、個人情報を本件業務の目的の範囲内でのみ利用し、漏洩・滅失・毀損の防止のため合理的な安全管理措置を講じる。",
                "3. 甲は、個人情報保護法上の個人情報取扱事業者であり、ユーザー・クリエイターに対する第一義的な責任を負う。",
            ],
        },
        {
            "heading": "第11条（本サービス運営に関する責任分担）",
            "items": [
                "1. 本サービスの運営者としての法的責任（利用規約・プライバシーポリシー・特定商取引法・プロバイダ責任制限法・資金決済法・個人情報保護法・その他関連法令に基づく責任）は、すべて甲が負う。乙は、甲の指示に基づき技術的実装・運用支援およびホスティング環境の提供を行う立場であり、本サービスの運営者としての責任を負わない。",
                "2. 第5条に基づき乙が自己名義でクラウドサービス等を契約・維持している場合であっても、本サービスおよび本サービスデータの運営主体・管理主体は甲であり、乙は技術的な実行環境を提供する受託者にすぎない。本サービスおよび本サービスデータに関して第三者（ユーザー・クリエイター・権利者・監督官庁等）から生じるあらゆる請求・苦情・訴訟・行政指導等は、甲が第一義的にこれに対応し、甲の責任および費用負担においてこれを解決する。",
                "3. ユーザー・クリエイターからの問合せ対応、投稿コンテンツの事前審査・事後審査・削除判断、違反者への措置、権利侵害申立対応、決済事業者・金融機関との折衝、個人情報保護法上の本人開示・訂正・削除請求への対応は、いずれも甲の責任および費用負担において行う。",
                "4. 乙は、甲の依頼に基づき、前項の業務に必要な技術的機能（管理画面・通知機能・ログ出力等）を実装する。",
                "5. 乙は、本サービスのシステム上、法令違反の疑いがあるコンテンツを発見した場合、速やかに甲に通知する。通知後の対応判断は甲が行う。",
                "6. 甲は、本サービスの利用規約、プライバシーポリシー、特定商取引法に基づく表記その他の法的文書において、本サービスの運営者が甲である旨を明記する。",
            ],
        },
        {
            "heading": "第12条（損害賠償）",
            "items": [
                "1. 甲および乙は、本契約の履行に関して、自己の故意または重大な過失により相手方に損害を与えた場合、当該損害を賠償する責任を負う。",
                "2. 乙が甲に対して負う損害賠償額の上限は、損害発生の直前12ヶ月間に甲が乙に支払った委託料（第4条の月額保守料および個別発注書に基づく委託料の合計）の金額を上限とする。ただし、乙の故意または重大な過失による場合はこの限りではない。",
                "3. 第三者（ユーザー・クリエイター・権利者等）から本サービスに関して甲に対して請求・訴訟が提起された場合、当該請求等が乙の故意または重大な過失に起因するものでない限り、甲の責任および費用において対応するものとし、乙に損害が生じた場合は甲がこれを補償する。",
            ],
        },
        {
            "heading": "第13条（反社会的勢力の排除）",
            "items": [
                "1. 甲および乙は、自己および自己の役員等が、暴力団、暴力団員、暴力団関係企業、総会屋、社会運動標榜ゴロその他これらに準ずる者（以下「反社会的勢力」という）に該当しないこと、および反社会的勢力と関係を有しないことを表明し、保証する。",
                "2. 前項に違反した場合、相手方は何らの催告なく本契約を解除することができる。",
            ],
        },
        {
            "heading": "第14条（契約期間）",
            "items": [
                "1. 本契約の有効期間は、2026年5月1日から1年間とする。",
                "2. 期間満了の1ヶ月前までに甲乙いずれからも書面による終了の意思表示がない場合、本契約は同一条件で1年間自動更新される。以後も同様とする。",
            ],
        },
        {
            "heading": "第15条（解除）",
            "items": [
                "1. 甲または乙は、相手方が本契約の条項に違反し、相当の期間を定めた催告にもかかわらず是正されない場合、本契約を解除することができる。",
                "2. 前項にかかわらず、相手方について以下のいずれかの事由が生じたときは、甲または乙は何らの催告なく本契約を解除することができる。",
                "　(1) 支払停止、支払不能、破産手続開始・民事再生手続開始・会社更生手続開始の申立てがあったとき",
                "　(2) 手形・小切手の不渡りを出したとき",
                "　(3) 差押・仮差押・仮処分・滞納処分を受けたとき",
                "　(4) 解散の決議をしたとき",
                "　(5) 第13条の規定に違反したとき",
            ],
        },
        {
            "heading": "第16条（契約終了時の措置）",
            "items": [
                "1. 本契約が終了した場合、乙は甲の指示に従い、以下の引継ぎ業務を行う。",
                "　(1) 成果物（ソースコード・設計書・運用ドキュメント等）の甲への引渡し",
                "　(2) 第5条に基づき乙名義で契約している本件ホスティング環境の、甲名義または甲が指定する第三者名義への移管支援",
                "　(3) 本サービスデータの甲への完全な引渡し",
                "　(4) 甲が指定する後任者への技術説明（合理的な範囲）",
                "2. 前項の引継ぎ業務にかかる費用については、甲乙協議のうえ決定する。",
                "3. 本件ホスティング環境の移管が完了するまでの間、乙は善良な管理者の注意をもって本件ホスティング環境を維持し、本サービスデータを保全する。この期間中のクラウドサービス利用料は、移管完了日までは乙が負担し、移管完了日以降は甲が負担する。",
                "4. 乙は、移管完了後、乙が保有する本サービスデータの複製を、甲の指示に従い速やかに削除する。ただし、法令により保存が義務付けられている場合、またはバックアップとして技術的に削除困難な場合はこの限りではない。",
            ],
        },
        {
            "heading": "第17条（残存条項）",
            "items": [
                "本契約終了後も、第7条（知的財産権）、第9条（秘密保持）、第10条（個人情報）、第11条（責任分担）、第12条（損害賠償）、第16条（契約終了時の措置）、第18条（準拠法・管轄）の規定は有効に存続する。",
            ],
        },
        {
            "heading": "第18条（準拠法・管轄）",
            "items": [
                "1. 本契約は日本法を準拠法とする。",
                "2. 本契約に関する一切の紛争については、東京地方裁判所を第一審の専属的合意管轄裁判所とする。",
            ],
        },
        {
            "heading": "第19条（協議事項）",
            "items": [
                "本契約に定めのない事項、または本契約の解釈に疑義が生じた場合は、甲乙誠実に協議のうえ解決する。",
            ],
        },
    ]

    for clause in clauses:
        add_h2(doc, clause["heading"])
        for item in clause["items"]:
            add_body(doc, item)

    doc.add_paragraph()
    p = doc.add_paragraph("本契約の成立を証するため、本書2通を作成し、甲乙各1通を保有する。")
    p.paragraph_format.line_spacing = 1.5

    doc.add_paragraph()
    date_p = doc.add_paragraph("　　　　年　　月　　日")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    party_table = doc.add_table(rows=4, cols=2)
    party_table.style = "Light Grid Accent 1"
    parties = [
        ("甲", "所在地", "東京都北区神谷2-21-7"),
        ("甲", "会社名", "株式会社91&Co."),
        ("甲", "代表者", "代表取締役　島瀬 直人　　　　　　　　　　　　　　　　㊞"),
    ]

    hdr = party_table.rows[0].cells
    hdr[0].text = "甲"
    hdr[1].text = ""
    set_cell_bg(hdr[0], "E8F0FE")
    set_cell_bg(hdr[1], "E8F0FE")
    for idx, (_, label, value) in enumerate(parties, start=1):
        party_table.rows[idx].cells[0].text = label
        party_table.rows[idx].cells[1].text = value

    doc.add_paragraph()

    party_table2 = doc.add_table(rows=4, cols=2)
    party_table2.style = "Light Grid Accent 1"
    parties2 = [
        ("所在地", "東京都渋谷区恵比寿西1-3-10 ファイブアネックス101"),
        ("会社名", "株式会社THINGx"),
        ("代表者", "代表取締役　植松 拓海　　　　　　　　　　　　　　　　㊞"),
    ]
    hdr2 = party_table2.rows[0].cells
    hdr2[0].text = "乙"
    hdr2[1].text = ""
    set_cell_bg(hdr2[0], "FEF0E8")
    set_cell_bg(hdr2[1], "FEF0E8")
    for idx, (label, value) in enumerate(parties2, start=1):
        party_table2.rows[idx].cells[0].text = label
        party_table2.rows[idx].cells[1].text = value

    out_path = os.path.join(HERE, "業務委託契約書_MyFocus.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")


def build_questions_doc() -> None:
    md_path = os.path.join(HERE, "naoto_questions.md")
    with open(md_path, encoding="utf-8") as f:
        md = f.read()

    doc = Document()
    set_default_font(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            doc.add_paragraph()
        elif line.startswith("# "):
            add_title(doc, line[2:].strip())
        elif line.startswith("## "):
            add_h2(doc, line[3:].strip())
        elif line.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(line[4:].strip())
            r.bold = True
            r.font.size = Pt(11)
        elif line.startswith("> "):
            add_body(doc, line[2:].strip())
        elif line.startswith("- [ ] "):
            add_list_item(doc, "☐ " + line[6:].strip())
        elif line.startswith("- "):
            add_list_item(doc, "・" + line[2:].strip())
        elif line.startswith("| "):
            # parse simple markdown table
            table_rows = []
            while i < len(lines) and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                table_rows.append(cells)
                i += 1
            # skip separator row if present
            cleaned = [r for r in table_rows if not all(set(c) <= set("-:") for c in r if c)]
            if cleaned:
                t = doc.add_table(rows=len(cleaned), cols=len(cleaned[0]))
                t.style = "Light Grid Accent 1"
                for ri, row in enumerate(cleaned):
                    for ci, val in enumerate(row):
                        t.rows[ri].cells[ci].text = val.replace("**", "")
                    if ri == 0:
                        for c in t.rows[0].cells:
                            set_cell_bg(c, "F2F2F2")
            continue
        elif line.startswith("---"):
            doc.add_paragraph()
        else:
            add_body(doc, line.replace("**", ""))
        i += 1

    out_path = os.path.join(HERE, "ナオト氏向け質問リスト_MyFocus.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build_contract_doc()
    build_questions_doc()
